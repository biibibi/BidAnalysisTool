#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多模态文档分析API模块
====================

本模块提供将包含图片和文字的docx文档直接作为大模型API输入的功能。
根据流程.txt的要求，实现文档内容提取、图片编码和多模态prompt构建。

主要功能：
    1. 从docx文档中提取所有文字内容和图片
    2. 将图片转换为base64编码
    3. 构建多模态prompt发送给大模型
    4. 支持Qwen-VL-Max等多模态模型

创建时间：2025年9月28日
"""

import os
import base64
import tempfile
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import json
import logging
from openai import OpenAI
from dotenv import load_dotenv

# 加载环境变量
load_dotenv(Path(__file__).resolve().parent / ".env", override=False)

# 配置日志
logger = logging.getLogger(__name__)


class MultimodalDocumentAnalyzer:
    """
    多模态文档分析器
    ================
    
    将包含图片和文字的Word文档直接转换为多模态API输入，
    实现真正意义上的"让AI理解完整的Word文档"。
    
    主要方法：
        - extract_text_and_images(): 提取文档内容
        - analyze_document(): 多模态文档分析
        - _build_multimodal_prompt(): 构建多模态提示词
    """
    
    def __init__(self):
        """
        初始化多模态文档分析器
        
        配置支持的大模型提供商和API参数
        """
        self.provider = os.getenv("LLM_PROVIDER", "qwen").lower()
        self._init_clients()
        
    def _init_clients(self):
        """初始化不同提供商的客户端"""
        # 初始化Qwen客户端
        self.qwen_client = None
        qwen_key = os.getenv("DASHSCOPE_API_KEY")
        if qwen_key:
            self.qwen_client = OpenAI(
                api_key=qwen_key,
                base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
            )
        
        # 初始化豆包客户端
        self.doubao_client = None
        ark_key = os.getenv("ARK_API_KEY")
        if ark_key:
            self.doubao_client = OpenAI(
                api_key=ark_key,
                base_url=os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
            )

    def extract_text_and_images(self, docx_path: Union[str, Path]) -> Dict[str, Any]:
        """
        从docx文档中提取文字内容和图片
        ==============================
        
        按照流程.txt的要求提取文档中的所有文字和图片内容。
        
        Args:
            docx_path: Word文档路径
            
        Returns:
            Dict: 包含文字内容和图片base64编码的字典
                - text_content: 文档的完整文字内容
                - images: 图片base64编码列表
                - image_count: 图片数量
                - success: 是否成功提取
                - error: 错误信息（如果有）
        """
        try:
            docx_path = Path(docx_path)
            if not docx_path.exists():
                return {
                    "success": False,
                    "error": f"文档不存在：{docx_path}",
                    "text_content": "",
                    "images": [],
                    "image_count": 0
                }
            
            # 加载文档
            doc = Document(docx_path)
            
            # 1. 提取所有文字内容
            text_lines = []
            
            # 提取段落文字
            for para in doc.paragraphs:
                if para.text.strip():
                    text_lines.append(para.text.strip())
            
            # 提取表格文字
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_lines.append(" | ".join(row_texts))
            
            full_text = "\n".join(text_lines)
            
            # 2. 提取和编码所有图片
            images_base64 = []
            
            # 创建临时目录保存图片
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                
                # 遍历文档中的图片关系
                image_counter = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        try:
                            # 获取图片二进制数据
                            image_data = rel.target_part.blob
                            
                            # 保存临时图片文件用于验证格式
                            img_temp_path = temp_path / f"temp_image_{image_counter}.png"
                            with open(img_temp_path, 'wb') as f:
                                f.write(image_data)
                            
                            # 使用PIL验证和标准化图片格式
                            with Image.open(BytesIO(image_data)) as img:
                                # 转换为RGB模式（确保兼容性）
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    img = img.convert('RGB')
                                
                                # 重新保存为标准格式
                                img_buffer = BytesIO()
                                img.save(img_buffer, format='JPEG', quality=85)
                                img_buffer.seek(0)
                                
                                # 转换为base64
                                encoded_img = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                                images_base64.append(f"data:image/jpeg;base64,{encoded_img}")
                                image_counter += 1
                                
                        except Exception as e:
                            logger.warning(f"处理图片时出错：{str(e)}")
                            continue
            
            logger.info(f"成功提取文档内容：文字 {len(full_text)} 字符，图片 {len(images_base64)} 张")
            
            return {
                "success": True,
                "text_content": full_text,
                "images": images_base64,
                "image_count": len(images_base64),
                "error": None
            }
            
        except Exception as e:
            error_msg = f"提取文档内容失败：{str(e)}"
            logger.error(error_msg)
            return {
                "success": False,
                "error": error_msg,
                "text_content": "",
                "images": [],
                "image_count": 0
            }

    def analyze_document(self, 
                        docx_path: Union[str, Path], 
                        analysis_prompt: str = None,
                        provider: str = None) -> Dict[str, Any]:
        """
        多模态文档分析
        ==============
        
        按照流程.txt的完整流程，将Word文档转换为多模态输入并调用大模型分析。
        
        Args:
            docx_path: Word文档路径
            analysis_prompt: 自定义分析提示词，如果不提供则使用默认的授权委托书分析
            provider: 指定使用的模型提供商，支持 'qwen' 或 'doubao'
            
        Returns:
            Dict: 分析结果
                - success: 是否成功
                - result: 大模型分析结果
                - text_content: 提取的文字内容
                - image_count: 图片数量
                - error: 错误信息（如果有）
        """
        try:
            # 1. 提取文档内容
            logger.info(f"开始分析文档：{docx_path}")
            extract_result = self.extract_text_and_images(docx_path)
            
            if not extract_result["success"]:
                return extract_result
            
            text_content = extract_result["text_content"]
            images = extract_result["images"]
            
            if not text_content and not images:
                return {
                    "success": False,
                    "error": "文档中没有找到文字内容或图片",
                    "result": None,
                    "text_content": "",
                    "image_count": 0
                }
            
            # 2. 构建多模态prompt
            multimodal_messages = self._build_multimodal_prompt(
                text_content, images, analysis_prompt
            )
            
            # 3. 调用多模态大模型API
            api_result = self._call_multimodal_api(multimodal_messages, provider)
            
            return {
                "success": api_result["success"],
                "result": api_result.get("result"),
                "text_content": text_content,
                "image_count": len(images),
                "error": api_result.get("error")
            }
            
        except Exception as e:
            error_msg = f"文档分析失败：{str(e)}"
            logger.error(error_msg)
            return {
                "success": False,
                "error": error_msg,
                "result": None,
                "text_content": "",
                "image_count": 0
            }

    def _build_multimodal_prompt(self, 
                                text_content: str, 
                                images: List[str], 
                                custom_prompt: str = None) -> List[Dict[str, Any]]:
        """
        构建多模态提示词
        ================
        
        按照流程.txt的格式要求构建多模态prompt。
        
        Args:
            text_content: 文档文字内容
            images: 图片base64编码列表
            custom_prompt: 自定义提示词
            
        Returns:
            List[Dict]: 多模态消息列表
        """
        # 构建消息内容列表
        content_parts = []
        
        # 添加任务说明
        content_parts.append({
            "type": "text",
            "text": custom_prompt
        })
        
        # 添加文字内容
        if text_content:
            content_parts.append({
                "type": "text", 
                "text": f"【文档文字内容】\n{text_content}"
            })
        
        # 添加图片内容
        if images:
            content_parts.append({
                "type": "text",
                "text": f"【文档包含图片】\n以下是文档中的 {len(images)} 张图片："
            })
            
            for i, img_base64 in enumerate(images, 1):
                content_parts.append({
                    "type": "text",
                    "text": f"图片 {i}："
                })
                content_parts.append({
                    "type": "image_url",
                    "image_url": {
                        "url": img_base64
                    }
                })
        
        # 添加最终分析要求
        content_parts.append({
            "type": "text",
            "text": "\n请综合文字内容和图片信息进行完整分析，确保信息的一致性和准确性。"
        })
        
        return [
            {
                "role": "user",
                "content": content_parts
            }
        ]

    def _call_multimodal_api(self, 
                           messages: List[Dict[str, Any]], 
                           provider: str = None) -> Dict[str, Any]:
        """
        调用多模态大模型API
        =================
        
        根据指定的提供商调用相应的多模态API。
        
        Args:
            messages: 多模态消息列表
            provider: 指定的提供商
            
        Returns:
            Dict: API调用结果
        """
        provider = provider or self.provider
        
        try:
            if provider in ("qwen", "ali", "dashscope"):
                return self._call_qwen_vl_api(messages)
            elif provider in ("doubao", "ark", "volc"):
                return self._call_doubao_vision_api(messages)
            else:
                # 默认使用Qwen
                return self._call_qwen_vl_api(messages)
                
        except Exception as e:
            return {
                "success": False,
                "error": f"API调用失败：{str(e)}",
                "result": None
            }

    def _call_qwen_vl_api(self, messages: List[Dict[str, Any]]) -> Dict[str, Any]:
        """调用Qwen-VL-Max API"""
        if not self.qwen_client:
            return {
                "success": False,
                "error": "Qwen客户端未初始化，请检查DASHSCOPE_API_KEY环境变量",
                "result": None
            }
        
        try:
            # 使用qwen3-vl-plus模型
            model = os.getenv("QWEN_VL_MODEL", "qwen3-vl-plus")
            
            response = self.qwen_client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.3,
                max_tokens=2000
            )
            
            result = response.choices[0].message.content
            
            logger.info(f"Qwen-VL API调用成功，返回 {len(result)} 字符")
            
            return {
                "success": True,
                "result": result,
                "error": None
            }
            
        except Exception as e:
            error_msg = f"Qwen-VL API调用失败：{str(e)}"
            logger.error(error_msg)
            return {
                "success": False,
                "error": error_msg,
                "result": None
            }

    def _call_doubao_vision_api(self, messages: List[Dict[str, Any]]) -> Dict[str, Any]:
        """调用豆包视觉API"""
        if not self.doubao_client:
            return {
                "success": False,
                "error": "豆包客户端未初始化，请检查ARK_API_KEY环境变量",
                "result": None
            }
        
        try:
            # 使用豆包视觉模型
            model = os.getenv("DOUBAO_VISION_MODEL", "doubao-vision")
            
            response = self.doubao_client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.3,
                max_tokens=2000
            )
            
            result = response.choices[0].message.content
            
            logger.info(f"豆包视觉API调用成功，返回 {len(result)} 字符")
            
            return {
                "success": True,
                "result": result,
                "error": None
            }
            
        except Exception as e:
            error_msg = f"豆包视觉API调用失败：{str(e)}"
            logger.error(error_msg)
            return {
                "success": False,
                "error": error_msg,
                "result": None
            }


# 便捷函数
def analyze_document_with_multimodal(docx_path: Union[str, Path], 
                                   analysis_prompt: str = None,
                                   provider: str = None) -> Dict[str, Any]:
    """
    便捷的多模态文档分析函数
    ========================
    
    这是一个快速使用的入口函数，按照流程.txt的要求实现完整的多模态文档分析。
    
    Args:
        docx_path: Word文档路径
        analysis_prompt: 自定义分析提示词
        provider: 指定模型提供商 ('qwen' 或 'doubao')
        
    Returns:
        Dict: 分析结果
        
    使用示例：
        # 分析授权委托书
        result = analyze_document_with_multimodal("授权委托书.docx")
        
        # 使用自定义提示词
        custom_prompt = "请分析这个投标文件的完整性..."
        result = analyze_document_with_multimodal("投标文件.docx", custom_prompt)
        
        # 指定使用豆包模型
        result = analyze_document_with_multimodal("文档.docx", provider="doubao")
    """
    analyzer = MultimodalDocumentAnalyzer()
    return analyzer.analyze_document(docx_path, analysis_prompt, provider)


def extract_document_content(docx_path: Union[str, Path]) -> Dict[str, Any]:
    """
    便捷的文档内容提取函数
    ======================
    
    只提取文档内容而不调用AI分析。
    
    Args:
        docx_path: Word文档路径
        
    Returns:
        Dict: 提取结果，包含text_content和images
    """
    analyzer = MultimodalDocumentAnalyzer()
    return analyzer.extract_text_and_images(docx_path)


# 主函数示例
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法：python multimodal_api.py <docx文件路径>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    print(f"开始分析文档：{docx_file}")
    print("=" * 50)
    
    # 执行多模态分析
    result = analyze_document_with_multimodal(docx_file)
    
    if result["success"]:
        print(f"✅ 分析成功！")
        print(f"📝 文字内容长度：{len(result['text_content'])} 字符")
        print(f"🖼️  图片数量：{result['image_count']} 张")
        print("\n📊 分析结果：")
        print("-" * 30)
        print(result["result"])
    else:
        print(f"❌ 分析失败：{result['error']}")