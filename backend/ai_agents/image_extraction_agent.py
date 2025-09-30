#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word 文档图片抽取 Agent（支持 AI 命名与选择性 OCR）"""

from __future__ import annotations

import json
import os
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from openai import OpenAI
from PIL import Image
from io import BytesIO
from dotenv import load_dotenv

from .base_agent import BaseAgent
from .ocr_agent import OCRAgent

try:
    import win32com.client  # type: ignore

    DOC_SUPPORT = True
except ImportError:  # pragma: no cover - win32 不可用时保持兼容
    DOC_SUPPORT = False

# 加载 backend/.env（若存在）
load_dotenv(Path(__file__).resolve().parent.parent / ".env", override=False)

AI_NAME_LOG = bool(os.getenv("AI_NAME_LOG"))

# ----------------------------- 数据模型 -----------------------------


@dataclass
class ImageContextSnapshot:
    """图片的上下文信息快照"""

    section_path: Tuple[str, ...] = field(default_factory=tuple)
    before_text: str = ""
    after_text: str = ""
    table_path: Optional[str] = None
    table_header: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    current_marker: Optional[str] = None

    def to_prompt_block(self) -> str:
        parts: List[str] = []
        if self.section_path:
            parts.append(f"章节路径：{' > '.join(self.section_path)}")
        if self.table_path:
            parts.append(f"表格位置：{self.table_path}")
        if self.table_header:
            parts.append(f"表格表头：{self.table_header}")
        if self.before_text:
            marker_line = f"\n{self.current_marker}" if self.current_marker else ""
            parts.append(f"图片前文：{self.before_text}{marker_line}")
        elif self.current_marker:
            parts.append(f"图片位置：{self.current_marker}")
        if self.after_text:
            parts.append(f"图片后文：{self.after_text}")
        return "\n".join(parts)


@dataclass
class ImageCandidate:
    """待处理的图片候选信息"""

    sequence_no: int
    embed_id: str
    image_part: Any
    block_index: int
    inline_index: int
    context: ImageContextSnapshot
    extension: str
    table_meta: Optional[Dict[str, Any]] = None
    group_id: Optional[str] = None
    group_index: int = 1
    group_size: int = 1
    order_index: int = 0


@dataclass
class SavedImageInfo:
    """落盘后的图片信息"""

    file_name: str
    file_path: str
    display_name: str
    naming_method: str
    ocr_status: str
    ocr_text: Optional[str]
    context: ImageContextSnapshot
    group_id: Optional[str]
    group_index: int
    group_size: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class NamingResult:
    base_name: str
    method: str
    dedupe_index: int


# ----------------------------- 工具方法 -----------------------------


def _is_heading(paragraph: Paragraph) -> bool:
    style = paragraph.style.name if paragraph.style else ""
    return bool(style and "heading" in style.lower())


def _heading_level(paragraph: Paragraph) -> int:
    style = paragraph.style.name if paragraph.style else ""
    match = re.search(r"(\d+)", style)
    return int(match.group(1)) if match else 1


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _limit_chars(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def _iter_block_items(parent: Any) -> Iterable[Any]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, Table):
        parent_elm = parent._tbl
    else:
        raise TypeError("Unsupported parent type for iterating blocks")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _collect_table_header(table: Table, max_rows: int = 2) -> str:
    headers: List[str] = []
    for row in table.rows[:max_rows]:
        row_text = []
        for cell in row.cells:
            text = _clean_text("\n".join(p.text for p in cell.paragraphs))
            if text:
                row_text.append(text)
        if row_text:
            headers.append(" | ".join(row_text))
    return "\n".join(headers)


def _detect_keywords(text: str) -> List[str]:
    KEYWORDS = [
        "身份证",
        "营业执照",
        "执照",
        "授权",
        "委托",
        "资质",
        "证书",
        "盖章",
        "签字",
        "合同",
        "架构图",
        "流程图",
    ]
    return [kw for kw in KEYWORDS if kw in text]


def _safe_mkdir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def _extract_image_extension(image_part: Any) -> str:
    try:
        with Image.open(BytesIO(image_part.blob)) as img:  # type: ignore[attr-defined]
            fmt = img.format.lower() if img.format else "png"
            if fmt == "jpeg":
                return "jpg"
            return fmt
    except Exception:
        return "png"


def _iter_run_fragments(run) -> Iterable[Tuple[str, str]]:  # type: ignore
    element = getattr(run, "_element", None)
    if element is None:
        return
    xml = getattr(element, "xml", None)
    if not xml:
        return

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }

    for child in run._element.iterchildren():
        tag = child.tag
        if tag.endswith("}t"):
            text = child.text or ""
            if text:
                yield ("text", text)
        elif tag.endswith("}tab"):
            yield ("text", "\t")
        elif tag.endswith("}cr") or tag.endswith("}br"):
            yield ("text", "\n")
        elif tag.endswith("}drawing"):
            blips = child.findall('.//a:blip', namespaces)
            for blip in blips:
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed:
                    yield ("image", embed)


# ----------------------------- AI 命名服务 -----------------------------


class ImageNamingService:
    def __init__(self) -> None:
        self.client = self._init_client()
        self.provider = os.getenv("LLM_PROVIDER", "qwen").lower()
        self.used_names: Dict[str, int] = {}

    def _init_client(self):
        provider = os.getenv("LLM_PROVIDER", "qwen").lower()
        try:
            if provider in ("qwen", "ali", "dashscope"):
                api_key = os.getenv("DASHSCOPE_API_KEY")
                if not api_key:
                    return None
                return OpenAI(
                    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
                    api_key=api_key,
                )
            if provider in ("doubao", "ark", "volc", "volcengine"):
                api_key = os.getenv("ARK_API_KEY")
                if not api_key:
                    return None
                base_url = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
                return OpenAI(base_url=base_url, api_key=api_key)
        except Exception as exc:  # pragma: no cover - 网络失败
            if AI_NAME_LOG:
                print(f"[ImageNamingService] 初始化LLM失败: {exc}")
        return None

    def generate_name(self, candidate: ImageCandidate) -> NamingResult:
        base_name = self._ai_name(candidate) if self.client else None
        method = "AI命名" if base_name else "规则命名"
        if not base_name:
            base_name = self._fallback_name(candidate)
        dedupe_index = self._register_name(base_name)
        return NamingResult(base_name=base_name, method=method, dedupe_index=dedupe_index)

    # ---------------- private helpers ----------------
    def _ai_name(self, candidate: ImageCandidate) -> Optional[str]:
        if not self.client:
            return None

        prompt = (
            "你作为资深文件档案管理员，需要为图片生成6-10个中文字符的文件名。"\
            "请根据以下上下文判断图片内容，并返回一个简洁的中文名（不含句号和空格）：\n\n"
            + candidate.context.to_prompt_block()
        )

        model = os.getenv("IMAGE_NAMING_MODEL")
        if not model:
            model = "qwen-turbo" if self.provider in ("qwen", "ali", "dashscope") else os.getenv("DOUBAO_MODEL_ID", "doubao-1.6-lite")

        try:
            response = self.client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "你是一个精准的命名助手，只返回图片文件名本身。",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.2,
                max_tokens=40,
                timeout=25,
            )
            name = response.choices[0].message.content.strip()
            name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]", "", name)
            if 2 <= len(name) <= 12:
                return name
        except Exception as exc:  # pragma: no cover - 网络调用
            if AI_NAME_LOG:
                print(f"[ImageNamingService] AI命名失败: {exc}")
        return None

    def _fallback_name(self, candidate: ImageCandidate) -> str:
        context = candidate.context
        # 优先使用章节末级标题
        if context.section_path:
            name = context.section_path[-1]
            name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]", "", name)
            if 2 <= len(name) <= 12:
                return name
        # 再尝试表格表头关键词
        if context.table_header:
            words = _clean_text(context.table_header).split(" | ")
            for word in words:
                clean = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]", "", word)
                if 2 <= len(clean) <= 12:
                    return clean
        # 关键词备选
        if context.keywords:
            return context.keywords[0]
        return "图片"

    def _register_name(self, base_name: str) -> int:
        counter = self.used_names.get(base_name, 0) + 1
        self.used_names[base_name] = counter
        return counter


# ----------------------------- Agent 实现 -----------------------------

class WordImageExtractionAgent(BaseAgent):
    """面向 Word 文档的图片抽取 Agent"""

    def __init__(self) -> None:
        super().__init__(
            name="WordImageExtractionAgent",
            description="提取 Word 图片并生成 AI 命名与可选 OCR 文本",
        )
        self.naming_service = ImageNamingService()
        self.ocr_agent = OCRAgent()

    # ------------------------------------------------------------------
    def process(self, content: str, context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        if not self.validate_input(content):
            return self.create_error_result("请输入有效的 Word 文件路径")

        try:
            file_path = Path(content).resolve()
            if not file_path.exists():
                return self.create_error_result(f"文件不存在: {file_path}")

            docx_path = self._ensure_docx(file_path)
            work_dir = self._resolve_work_dir(docx_path, context)
            images_dir = _safe_mkdir(work_dir / "images")

            document = Document(docx_path)
            image_parts = {
                rel.rId: rel.target_part
                for rel in document.part.rels.values()
                if "image" in rel.reltype
            }

            text_blocks, candidates = self._collect_candidates(document, image_parts)
            self._attach_context(text_blocks, candidates)
            self._assign_groups(candidates)

            saved = self._save_images(images_dir, candidates)
            manifest_path = self._write_manifest(work_dir, saved, docx_path)

            data = {
                "image_count": len(saved),
                "work_dir": str(work_dir),
                "images_dir": str(images_dir),
                "manifest_path": str(manifest_path) if manifest_path else None,
                "images": [self._serialize_saved_image(item) for item in saved],
            }
            return self.create_success_result(data, "图片提取完成")

        except Exception as exc:
            self.logger.exception("图片抽取失败")
            return self.create_error_result(f"图片抽取失败: {exc}")

    # ------------------------------------------------------------------
    def _ensure_docx(self, file_path: Path) -> Path:
        if file_path.suffix.lower() == ".docx":
            return file_path
        if file_path.suffix.lower() != ".doc":
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
        if not DOC_SUPPORT:
            raise RuntimeError("缺少 .doc 转换环境，请安装 pywin32")

        target = file_path.with_suffix(".docx")
        word = win32com.client.Dispatch("Word.Application")  # type: ignore
        word.Visible = False
        try:
            doc = word.Documents.Open(str(file_path))
            doc.SaveAs(str(target), FileFormat=16)
            doc.Close()
        finally:
            word.Quit()
        return target

    def _resolve_work_dir(self, docx_path: Path, context: Optional[Dict[str, Any]]) -> Path:
        if context and context.get("work_dir"):
            return _safe_mkdir(Path(context["work_dir"]))
        default_dir = docx_path.parent / f"{docx_path.stem}_assets"
        return _safe_mkdir(default_dir)

    # ------------------------------------------------------------------
    def _collect_candidates(
        self, document: Document, image_parts: Dict[str, Any]
    ) -> Tuple[List[Dict[str, Any]], List[ImageCandidate]]:
        text_blocks: List[Dict[str, Any]] = []
        candidates: List[ImageCandidate] = []
        section_stack: List[str] = []
        block_index = -1
        sequence_no = 0
        order_counter = 0
        table_counter = 0

        def append_text_block(
            block_idx: int,
            text: str,
            section: Tuple[str, ...],
            table_path: Optional[str],
            *,
            is_heading: bool = False,
        ) -> None:
            nonlocal order_counter
            cleaned = _clean_text(text)
            if not cleaned:
                return
            text_blocks.append(
                {
                    "index": block_idx,
                    "order": order_counter,
                    "text": cleaned,
                    "section": section,
                    "table_path": table_path,
                    "is_heading": is_heading,
                }
            )
            order_counter += 1

        def append_placeholder(
            block_idx: int, section: Tuple[str, ...], table_path: Optional[str], seq: int
        ) -> None:
            nonlocal order_counter
            text_blocks.append(
                {
                    "index": block_idx,
                    "order": order_counter,
                    "text": f"[图片{seq}]",
                    "section": section,
                    "table_path": table_path,
                    "placeholder": True,
                }
            )
            order_counter += 1

        for block in _iter_block_items(document):
            block_index += 1
            if isinstance(block, Paragraph):
                full_text = _clean_text(block.text)
                if _is_heading(block) and full_text:
                    level = _heading_level(block)
                    section_stack = section_stack[: level - 1] + [full_text]

                is_heading_para = _is_heading(block)
                inline_position = 0
                for run in block.runs:
                    for frag_type, value in _iter_run_fragments(run):
                        if frag_type == "text":
                            append_text_block(
                                block_index,
                                value,
                                tuple(section_stack),
                                None,
                                is_heading=is_heading_para,
                            )
                            inline_position += 1
                        elif frag_type == "image":
                            embed_id = value
                            if embed_id not in image_parts:
                                continue
                            sequence_no += 1
                            candidate = ImageCandidate(
                                sequence_no=sequence_no,
                                embed_id=embed_id,
                                image_part=image_parts[embed_id],
                                block_index=block_index,
                                inline_index=inline_position,
                                context=ImageContextSnapshot(section_path=tuple(section_stack)),
                                extension=_extract_image_extension(image_parts[embed_id]),
                                order_index=order_counter,
                            )
                            candidate.context.current_marker = f"[图片{sequence_no}]"
                            candidates.append(candidate)
                            append_placeholder(block_index, tuple(section_stack), None, sequence_no)
                            inline_position += 1
            elif isinstance(block, Table):
                table_counter += 1
                table_idx = table_counter
                header_text = _collect_table_header(block)
                for row_idx, row in enumerate(block.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_table_path = f"表格{table_idx} 行{row_idx+1} 列{col_idx+1}"
                        inline_position = 0
                        for para in cell.paragraphs:
                            for run in para.runs:
                                for frag_type, value in _iter_run_fragments(run):
                                    if frag_type == "text":
                                        append_text_block(
                                            block_index,
                                            value,
                                            tuple(section_stack),
                                            cell_table_path,
                                        )
                                        inline_position += 1
                                    elif frag_type == "image":
                                        embed_id = value
                                        if embed_id not in image_parts:
                                            continue
                                        sequence_no += 1
                                        table_meta = {
                                            "table_index": table_idx,
                                            "row_index": row_idx + 1,
                                            "column_index": col_idx + 1,
                                            "header": header_text,
                                        }
                                        candidate = ImageCandidate(
                                            sequence_no=sequence_no,
                                            embed_id=embed_id,
                                            image_part=image_parts[embed_id],
                                            block_index=block_index,
                                            inline_index=inline_position,
                                            context=ImageContextSnapshot(
                                                section_path=tuple(section_stack),
                                                table_path=cell_table_path,
                                                table_header=_limit_chars(header_text, 120) if header_text else None,
                                            ),
                                            extension=_extract_image_extension(image_parts[embed_id]),
                                            table_meta=table_meta,
                                            order_index=order_counter,
                                        )
                                        candidate.context.current_marker = f"[图片{sequence_no}]"
                                        candidates.append(candidate)
                                        append_placeholder(block_index, tuple(section_stack), cell_table_path, sequence_no)
                                        inline_position += 1
        return text_blocks, candidates

    def _attach_context(self, text_blocks: List[Dict[str, Any]], candidates: List[ImageCandidate]) -> None:
        if not candidates:
            return
        ordered_blocks = sorted(text_blocks, key=lambda x: x.get("order", 0))
        for candidate in candidates:
            relevant = [
                item
                for item in ordered_blocks
                if item.get("order", 0) < candidate.order_index and item.get("text")
            ]
            collected: List[Dict[str, Any]] = []
            for item in reversed(relevant):
                if item.get("placeholder"):
                    break
                collected.append(item)
                if item.get("is_heading"):
                    break

            before_fragments = [entry["text"] for entry in reversed(collected)]
            before_text = " \n".join(before_fragments).strip()

            candidate.context.before_text = before_text
            candidate.context.after_text = ""

            combined_text = " ".join([before_text, candidate.context.table_header or ""]).strip()
            candidate.context.keywords = _detect_keywords(combined_text)

            marker = candidate.context.current_marker or f"[图片{candidate.sequence_no}]"
            candidate.context.current_marker = marker

    # ------------------------------------------------------------------
    def _assign_groups(self, candidates: List[ImageCandidate]) -> None:
        if not candidates:
            return
        candidates.sort(key=lambda c: c.sequence_no)
        group_id = None
        current_group: List[ImageCandidate] = []

        def flush_group():
            if not current_group:
                return
            if len(current_group) == 1:
                current_group[0].group_id = None
                current_group[0].group_size = 1
                current_group[0].group_index = 1
                return
            gid = str(uuid.uuid4())
            for idx, cand in enumerate(current_group, start=1):
                cand.group_id = gid
                cand.group_index = idx
                cand.group_size = len(current_group)

        prev = None
        for candidate in candidates:
            if prev is None:
                current_group = [candidate]
                prev = candidate
                continue
            same_block = candidate.block_index == prev.block_index
            close_block = abs(candidate.block_index - prev.block_index) <= 1
            same_table = (candidate.context.table_path and prev.context.table_path and candidate.context.table_path.split(" 行")[0] == prev.context.table_path.split(" 行")[0])
            if same_block or (close_block and (same_table or not candidate.context.table_path and not prev.context.table_path)):
                current_group.append(candidate)
            else:
                flush_group()
                current_group = [candidate]
            prev = candidate
        flush_group()

    # ------------------------------------------------------------------
    def _save_images(self, images_dir: Path, candidates: List[ImageCandidate]) -> List[SavedImageInfo]:
        saved: List[SavedImageInfo] = []
        if not candidates:
            return saved

        ocr_mode = os.getenv("IMAGE_OCR_MODE", "never").lower()
        group_display_map: Dict[str, Dict[str, Any]] = {}
        used_display_names: Dict[str, int] = {}

        for candidate in candidates:
            naming_result = self.naming_service.generate_name(candidate)
            base_name = naming_result.base_name
            if naming_result.dedupe_index > 1:
                base_name = f"{base_name}_{naming_result.dedupe_index}"

            sequence_prefix = str(candidate.sequence_no).zfill(3)
            display_name: str
            naming_method: str

            if candidate.group_id:
                info = group_display_map.get(candidate.group_id)
                if not info:
                    info = {
                        "base": base_name,
                        "method": naming_result.method,
                    }
                    group_display_map[candidate.group_id] = info
                base_for_group = info["base"]
                naming_method = info["method"]
                display_name = f"{sequence_prefix}_{base_for_group}"
                if candidate.group_size > 1 and candidate.group_index > 1:
                    display_name = f"{display_name}_{candidate.group_index}"
            else:
                display_name = f"{sequence_prefix}_{base_name}"
                naming_method = naming_result.method

            # 确保文件名唯一
            display_base = display_name
            counter = used_display_names.get(display_base, 0)
            if counter:
                display_name = f"{display_base}_{counter+1}"
            used_display_names[display_base] = counter + 1

            file_name = f"{display_name}.{candidate.extension}"
            file_path = images_dir / file_name
            counter = 1
            while file_path.exists():
                file_name = f"{display_name}_{counter}.{candidate.extension}"
                file_path = images_dir / file_name
                counter += 1

            with open(file_path, "wb") as f:
                f.write(candidate.image_part.blob)  # type: ignore[attr-defined]

            ocr_status, ocr_text = self._maybe_run_ocr(candidate, file_path, mode=ocr_mode)

            saved.append(
                SavedImageInfo(
                    file_name=file_name,
                    file_path=str(file_path),
                    display_name=display_name,
                    naming_method=naming_method,
                    ocr_status=ocr_status,
                    ocr_text=ocr_text,
                    context=candidate.context,
                    group_id=candidate.group_id,
                    group_index=candidate.group_index,
                    group_size=candidate.group_size,
                    metadata={
                        "sequence_no": candidate.sequence_no,
                        "table_meta": candidate.table_meta,
                        "placeholder": f"[图片{candidate.sequence_no}]",
                        "order_index": candidate.order_index,
                    },
                )
            )
        return saved

    def _maybe_run_ocr(self, candidate: ImageCandidate, file_path: Path, mode: str = "auto") -> Tuple[str, Optional[str]]:
        mode = mode or "auto"
        mode = mode.lower()
        if mode == "never":
            return "skipped", None
        if mode == "always":
            return self._run_ocr(file_path)
        if mode == "auto":
            if self._should_run_ocr(candidate):
                return self._run_ocr(file_path)
            return "skipped", None
        return "skipped", None

    def _should_run_ocr(self, candidate: ImageCandidate) -> bool:
        keywords = set(candidate.context.keywords)
        preferred = {"身份证", "营业执照", "证书", "授权", "合同"}
        if keywords & preferred:
            return True
        # 表格类且表头含关键字
        header = candidate.context.table_header or ""
        if any(word in header for word in preferred):
            return True
        # 如果章节标题明确包含关键词
        if any(word in "".join(candidate.context.section_path) for word in preferred):
            return True
        return False

    def _run_ocr(self, file_path: Path) -> Tuple[str, Optional[str]]:
        try:
            result = self.ocr_agent.general_ocr(str(file_path))
            if result.get("success"):
                text = result.get("data", {}).get("text")
                return "processed", _limit_chars(text or "", 400)
            return "failed", None
        except Exception as exc:  # pragma: no cover - OCR 失败不应中断主流程
            if AI_NAME_LOG:
                print(f"[OCR] 调用失败: {exc}")
            return "failed", None

    # ------------------------------------------------------------------
    def _write_manifest(self, work_dir: Path, saved: Sequence[SavedImageInfo], docx_path: Path) -> Optional[Path]:
        if not saved:
            return None
        manifest = {
            "document": str(docx_path),
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "image_count": len(saved),
            "images": [
                {
                    "file_name": item.file_name,
                    "file_path": item.file_path,
                    "display_name": item.display_name,
                    "naming_method": item.naming_method,
                    "ocr_status": item.ocr_status,
                    "ocr_text": item.ocr_text,
                    "context": {
                        "section_path": list(item.context.section_path),
                        "before_text": item.context.before_text,
                        "after_text": item.context.after_text,
                        "table_path": item.context.table_path,
                        "table_header": item.context.table_header,
                        "keywords": item.context.keywords,
                    },
                    "group": {
                        "id": item.group_id,
                        "index": item.group_index,
                        "size": item.group_size,
                    },
                    "metadata": item.metadata,
                }
                for item in saved
            ],
        }
        manifest_path = work_dir / "images_manifest.json"
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        return manifest_path

    def _serialize_saved_image(self, item: SavedImageInfo) -> Dict[str, Any]:
        return {
            "file_name": item.file_name,
            "file_path": item.file_path,
            "display_name": item.display_name,
            "naming_method": item.naming_method,
            "ocr_status": item.ocr_status,
            "ocr_text": item.ocr_text,
            "context": {
                "section_path": list(item.context.section_path),
                "before_text": item.context.before_text,
                "after_text": item.context.after_text,
                "table_path": item.context.table_path,
                "table_header": item.context.table_header,
                "keywords": item.context.keywords,
            },
            "group": {
                "id": item.group_id,
                "index": item.group_index,
                "size": item.group_size,
            },
            "metadata": item.metadata,
        }


def create_image_extraction_agent() -> WordImageExtractionAgent:
    return WordImageExtractionAgent()
