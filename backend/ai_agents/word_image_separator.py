#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档图片分离工具 - 简化版
输入: Word文档 (.docx/.doc)
输出: 分离后的Word文档 + 图片文件夹
"""

import os
import sys
from pathlib import Path
import re

try:
    from docx import Document
    from PIL import Image
    from io import BytesIO
    from openai import OpenAI
    from dotenv import load_dotenv
except ImportError as e:
    print(f"缺少依赖库: {e}")
    print("请安装: pip install python-docx pillow openai")
    sys.exit(1)

# 可选的.doc文件支持
try:
    import win32com.client
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False
    print("提示: 无法处理.doc文件 (需要: pip install pywin32)")

# 加载环境变量
load_dotenv(Path(__file__).parent.parent / '.env')

# 初始化火山大模型客户端
def init_ai_client():
    """初始化AI客户端"""
    try:
        # 读取环境变量
        ark_api_key = os.getenv('ARK_API_KEY')
        if not ark_api_key:
            print("警告: 未设置ARK_API_KEY环境变量")
            return None
            
        client = OpenAI(
            base_url="https://ark.cn-beijing.volces.com/api/v3",
            api_key=ark_api_key,
        )
        return client
    except Exception as e:
        print(f"警告: 无法初始化AI客户端: {e}")
        return None


def generate_image_name(ai_client, context_text, image_index=1, total_images=1):
    """
    使用AI生成图片名称
    
    Args:
        ai_client: AI客户端
        context_text: 图片周围的文字内容
        image_index: 当前图片在连续图片中的索引 
        total_images: 连续图片总数
        
    Returns:
        str: 生成的图片名称（不含扩展名）
    """
    if not ai_client:
        return f"图片_{image_index:03d}"
    
    try:
        # 构建提示词 - 针对身份证明文档优化
        if total_images > 1:
            prompt = f"""
你是专业的文档分析师，需要根据上下文准确判断图片内容并生成合适的文件名。

上下文信息：
{context_text[:800]}

分析要求：
1. 仔细阅读上下文，重点关注距离图片最近的标题和描述
2. 根据最相关的上下文内容判断图片类型
3. 生成准确的4-8个中文字符的文件名
4. 这是第{image_index}张图片，共{total_images}张连续图片

常见图片类型参考：
- 技术方案图、系统架构图、功能模块图
- GPS管理图、监控管理图、数据管理图
- 身份证、执照、证书、授权书
- 项目组成员、投标单位信息
- 其他相关文档图片

注意事项：
- 优先使用上下文中出现的具体名词
- 如果上下文提到"GPS管理"，命名为"GPS管理"而不是其他
- 如果上下文提到"功能场景"，命名为"功能场景"而不是其他
- 只返回文件名，不要任何解释或额外文字
- 避免使用特殊字符

请直接返回最合适的文件名："""
        else:
            prompt = f"""
你是专业的文件命名助手。请基于以下上下文为图片生成一个准确的中文文件名。

上下文内容：
{context_text[:800]}

特别说明：
- 技术方案图、系统架构图、功能模块图
- GPS管理图、监控管理图、数据管理图
- 身份证、执照、证书、授权书
- 项目组成员、投标单位信息

要求：
1. 根据上下文判断图片类型，优先使用上下文中的关键词
2. 使用4-8个中文字符  
3. 只返回文件名，不要其他解释
4. 避免使用特殊字符和空格

示例：功能场景建设、GPS管理系统、法定代表人身份证
"""

        # 获取豆包模型ID
        model_id = os.getenv('DOUBAO_MODEL_ID', 'doubao-seed-1-6-250615')
        
        # 调试日志
        if os.getenv('AI_NAME_LOG'):
            print(f"🤖 使用豆包模型: {model_id}")
            print(f"📝 上下文长度: {len(context_text)}")
            print(f"📄 上下文内容: {context_text[:100]}...")
        
        completion = ai_client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": "你是一个专业的文件命名助手，擅长根据文档内容生成合适的图片文件名。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=50,
            timeout=10  # 添加10秒超时
        )
        
        generated_name = completion.choices[0].message.content.strip()
        
        # 调试日志
        if os.getenv('AI_NAME_LOG'):
            print(f"🎯 AI生成名称: {generated_name}")
        
        # 清理文件名，移除特殊字符
        cleaned_name = re.sub(r'[^\w\u4e00-\u9fff]', '', generated_name)
        
        # 如果是连续图片且生成的名称没有数字后缀，添加后缀
        if total_images > 1 and not re.search(r'\d+$', cleaned_name):
            cleaned_name = f"{cleaned_name}_{image_index}"
            
        return cleaned_name if cleaned_name else f"图片_{image_index:03d}"
        
    except Exception as e:
        print(f"AI命名失败: {e}")
        # 简单基于上下文生成描述性名称
        if context_text:
            # 提取关键词
            words = context_text.strip()[:50]
            if "认证" in words:
                return f"认证_{image_index:03d}"
            elif "证书" in words:
                return f"证书_{image_index:03d}"
            elif "测试" in words or "检测" in words:
                return f"测试_{image_index:03d}"
            elif "执照" in words or "营业" in words:
                return f"执照_{image_index:03d}"
            elif "授权" in words:
                return f"授权_{image_index:03d}"
        
        return f"图片_{image_index:03d}"


def is_empty_or_whitespace_paragraph(paragraph):
    """判断段落是否为空或只包含空白符"""
    text = paragraph.text.strip()
    return len(text) == 0


def collect_headings(doc):
    """收集文档中的标题信息"""
    headings = []
    for i, para in enumerate(doc.paragraphs):
        if para.style and ('Heading' in para.style.name or 'heading' in para.style.name.lower()):
            text = para.text.strip()
            if text:
                headings.append({
                    'index': i,
                    'text': text,
                    'style': para.style.name,
                    'level': getattr(para.style, 'level', 1) or 1
                })
    return headings


def find_section_for_image(headings, image_para_idx, total_paras):
    """找到图片所属的章节"""
    current_section = None
    next_section_start = total_paras
    
    for i, heading in enumerate(headings):
        if heading['index'] < image_para_idx:
            current_section = heading
        elif heading['index'] > image_para_idx:
            next_section_start = heading['index']
            break
    
    if current_section:
        section_start = current_section['index']
        section_end = next_section_start
        return current_section, section_start, section_end
    
    return None, 0, total_paras


def extract_smart_name_from_contexts(contexts):
    """
    从上下文列表中智能提取图片名称
    
    Args:
        contexts: 上下文段落列表
        
    Returns:
        str: 提取的图片名称，如果无法提取则返回None
    """
    if not contexts:
        return None
        
    # 定义不同类型的关键词模式和对应的名称
    naming_patterns = [
        # 证件类
        {
            'keywords': ['营业执照', '执照'],
            'name': '营业执照'
        },
        {
            'keywords': ['身份证'],
            'name': '身份证'
        },
        {
            'keywords': ['法定代表人', '法人代表'],
            'name': '法定代表人身份证'
        },
        {
            'keywords': ['授权委托书', '委托书'],
            'name': '授权委托书'
        },
        {
            'keywords': ['授权代理人', '代理人'],
            'name': '授权代理人身份证'
        },
        
        # 技术方案类
        {
            'keywords': ['功能场景建设', '功能场景'],
            'name': '功能场景建设'
        },
        {
            'keywords': ['技术架构设计', '技术架构'],
            'name': '技术架构设计'
        },
        {
            'keywords': ['总体架构设计', '总体架构'],
            'name': '总体架构设计'
        },
        {
            'keywords': ['系统架构', '架构设计'],
            'name': '系统架构'
        },
        
        # GPS和管理类
        {
            'keywords': ['GPS管理', 'GPS系统'],
            'name': 'GPS管理'
        },
        {
            'keywords': ['项目GPS管理'],
            'name': '项目GPS管理'
        },
        {
            'keywords': ['车辆GPS明细', '车辆GPS'],
            'name': '车辆GPS明细'
        },
        {
            'keywords': ['项目设备管理', '设备管理'],
            'name': '项目设备管理'
        },
        {
            'keywords': ['视频监控系统监控管理', '视频监控管理', '监控管理'],
            'name': '监控管理'
        },
        {
            'keywords': ['轨迹点位管理', '点位管理'],
            'name': '点位管理'
        },
        {
            'keywords': ['新增点位'],
            'name': '新增点位'
        },
        {
            'keywords': ['回放工作台', '回放'],
            'name': '回放工作台'
        },
        {
            'keywords': ['智能筛选'],
            'name': '智能筛选'
        },
        {
            'keywords': ['告警管理'],
            'name': '告警管理'
        },
        {
            'keywords': ['GPS异常'],
            'name': 'GPS异常'
        },
        {
            'keywords': ['GPS轨迹偏移异常', '轨迹偏移'],
            'name': 'GPS轨迹偏移'
        },
        {
            'keywords': ['车辆异常停留'],
            'name': '车辆异常停留'
        },
        {
            'keywords': ['非工作时间异常'],
            'name': '非工作时间异常'
        },
        {
            'keywords': ['巡查管理'],
            'name': '巡查管理'
        },
        {
            'keywords': ['综合驾驶舱'],
            'name': '综合驾驶舱'
        },
        
        # 投标相关
        {
            'keywords': ['投标单位实力', '单位实力'],
            'name': '投标单位实力'
        },
        {
            'keywords': ['项目组成员'],
            'name': '项目组成员'
        },
        {
            'keywords': ['技术和商务符合性', '符合性'],
            'name': '技术商务符合性'
        }
    ]
    
    # 检查每个上下文段落，找到最匹配的模式
    for context in contexts:
        text = context['text']
        
        # 按优先级检查模式（优先匹配更具体的关键词）
        for pattern in naming_patterns:
            for keyword in pattern['keywords']:
                if keyword in text:
                    return pattern['name']
    
    # 如果没有匹配到预定义模式，尝试从标题中提取
    for context in contexts:
        if context.get('is_heading') and context['text']:
            title = context['text'].strip()
            # 清理标题，保留主要部分
            clean_title = re.sub(r'[^\w\u4e00-\u9fff]', '', title)
            if 4 <= len(clean_title) <= 12:
                return clean_title
    
    return None


def find_nearest_context_paragraphs(doc, image_para_idx, max_distance=20):
    """
    为图片寻找最近的5段有效上下文文字
    
    Args:
        doc: Word文档对象
        image_para_idx: 图片所在段落索引
        max_distance: 最大搜索距离
        
    Returns:
        list: 最有效的5段上下文段落信息
    """
    context_candidates = []
    
    # 向前和向后搜索有意义的内容
    search_range = range(
        max(0, image_para_idx - max_distance), 
        min(len(doc.paragraphs), image_para_idx + max_distance + 1)
    )
    
    for para_idx in search_range:
        if para_idx >= len(doc.paragraphs) or para_idx == image_para_idx:
            continue
            
        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        
        if not text:
            continue
            
        # 检查是否是标题
        is_heading = para.style and ('Heading' in para.style.name or 'heading' in para.style.name.lower())
        
        # 过滤掉完全无意义的内容
        ignore_patterns = [
            '[图片:', '图片说明', '如下图', '见下图', '上图', '下图',
            '图所示', '如图', '插图', '附图'
        ]
        
        # 超高价值的描述性内容（证件类）
        super_high_value_patterns = [
            '身份证', '营业执照', '法定代表人', '授权代理人', '授权委托人',
            '单位负责人', '自然人', '组织机构代码证', '税务登记证',
            '开户许可证', '银行账户', '资质证书'
        ]
        
        # 高价值的描述性内容
        high_value_patterns = [
            '扫描件', '复印件', '正本', '副本', '原件', '附件',
            '证书', '认证', '许可证', '证明', '委托书',
            '资质', '执照', '登记', '备案', '审批'
        ]
        
        # 中等价值的内容
        medium_value_patterns = [
            '公司', '企业', '单位', '姓名', '名称', '地址', 
            '电话', '联系', '职务', '部门', '说明', '材料',
            '投标人', '供应商', '承包商', '甲方', '乙方'
        ]
        
        # 低价值但仍有用的内容
        low_value_patterns = [
            '项目', '工程', '采购', '招标', '投标', '合同',
            '服务', '产品', '技术', '方案', '要求'
        ]
        
        should_ignore = False
        for pattern in ignore_patterns:
            if pattern in text:
                should_ignore = True
                break
                
        if should_ignore:
            continue
        
        # 计算段落价值分数（更精确的评估）
        distance = abs(para_idx - image_para_idx)
        value_score = 0
        
        # 基础分数：距离越近分数越高（非线性衰减）
        if distance <= 3:
            distance_score = 30 - (distance * 5)  # 距离1-3段: 25-15分
        elif distance <= 8:
            distance_score = 15 - (distance - 3) * 2  # 距离4-8段: 13-3分
        else:
            distance_score = max(1, 10 - distance)  # 距离9+段: 递减到1分
        
        value_score += distance_score
        
        # 标题特别加分
        if is_heading:
            value_score += 40
            
        # 内容价值加分（分层评估）
        content_bonus = 0
        for pattern in super_high_value_patterns:
            if pattern in text:
                content_bonus = max(content_bonus, 50)  # 超高价值
                
        if content_bonus == 0:  # 如果没有超高价值，检查高价值
            for pattern in high_value_patterns:
                if pattern in text:
                    content_bonus = max(content_bonus, 30)  # 高价值
                    
        if content_bonus == 0:  # 如果没有高价值，检查中等价值
            for pattern in medium_value_patterns:
                if pattern in text:
                    content_bonus = max(content_bonus, 15)  # 中等价值
                    
        if content_bonus == 0:  # 如果没有中等价值，检查低价值
            for pattern in low_value_patterns:
                if pattern in text:
                    content_bonus = max(content_bonus, 5)  # 低价值
        
        value_score += content_bonus
        
        # 长度评估（信息密度）
        text_length = len(text)
        if 10 <= text_length <= 50:  # 最佳长度
            value_score += 10
        elif 5 <= text_length <= 100:  # 良好长度
            value_score += 5
        elif text_length > 200:  # 太长扣分
            value_score -= 5
            
        # 检查是否包含关键信息
        has_key_info = any(keyword in text for keyword in super_high_value_patterns + high_value_patterns)
        
        # 位置偏好（图片前后的内容可能更相关）
        if para_idx < image_para_idx:  # 图片前的内容
            position_bonus = 3
        else:  # 图片后的内容
            position_bonus = 2
        value_score += position_bonus
        
        context_candidates.append({
            'para_idx': para_idx,
            'text': text,
            'distance': distance,
            'value_score': value_score,
            'is_heading': is_heading,
            'has_key_info': has_key_info,
            'text_length': text_length,
            'content_bonus': content_bonus,
            'distance_score': distance_score
        })
    
    # 按价值分数排序，选择最好的5个
    context_candidates.sort(key=lambda x: x['value_score'], reverse=True)
    
    # 返回最有效的5个段落
    return context_candidates[:5]


def find_nearby_meaningful_content(doc, image_para_idx, max_distance=10):
    """
    查找图片附近的有意义标题或段落内容
    
    Args:
        doc: Word文档对象
        image_para_idx: 图片所在段落索引
        max_distance: 最大搜索距离
        
    Returns:
        tuple: (content_text, content_type, context_paragraphs) 内容文本、类型和上下文段落
    """
    meaningful_content = []
    context_paragraphs = []
    
    # 向前和向后搜索有意义的内容
    search_range = range(
        max(0, image_para_idx - max_distance), 
        min(len(doc.paragraphs), image_para_idx + max_distance + 1)
    )
    
    for para_idx in search_range:
        if para_idx >= len(doc.paragraphs):
            continue
            
        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        
        if not text:
            continue
            
        # 检查是否是标题
        is_heading = para.style and ('Heading' in para.style.name or 'heading' in para.style.name.lower())
        
        # 过滤掉无意义的内容，但保留描述性内容
        ignore_patterns = [
            '[图片:', '签名', '盖章', '电子签名'
        ]
        
        # 特别关注的描述性内容
        descriptive_patterns = [
            '扫描件', '复印件', '正本', '副本', '原件', '附件',
            '身份证', '营业执照', '执照', '证书', '认证', 
            '授权代理人', '法定代表人', '单位负责人',
            '资质', '许可证', '证明', '委托书'
        ]
        
        is_meaningful = True
        is_descriptive = False
        
        # 检查是否包含描述性内容
        for pattern in descriptive_patterns:
            if pattern in text:
                is_descriptive = True
                break
                
        # 检查是否应该忽略
        for pattern in ignore_patterns:
            if pattern in text:
                is_meaningful = False
                break
        
        # 如果文字太短但包含描述性内容，也认为是有意义的
        if is_meaningful and (len(text) > 3 or is_descriptive):
            distance = abs(para_idx - image_para_idx)
            
            # 计算优先级：标题 > 描述性内容 > 普通内容
            if is_heading:
                priority = 1000
            elif is_descriptive:
                priority = 500  # 描述性内容优先级较高
            else:
                priority = 100
                
            priority -= distance  # 距离越近优先级越高
            
            # 提取关键词
            key_words = []
            content_keywords = [
                '身份证', '营业执照', '执照', '证书', '认证', 
                '授权代理人', '法定代表人', '单位负责人',
                '资质', '许可证', '证明', '委托书'
            ]
            
            for keyword in content_keywords:
                if keyword in text:
                    key_words.append(keyword)
            
            content_item = {
                'text': text,
                'para_idx': para_idx,
                'distance': distance,
                'is_heading': is_heading,
                'is_descriptive': is_descriptive,
                'priority': priority,
                'keywords': key_words
            }
            
            meaningful_content.append(content_item)
            
            # 收集有意义的段落作为上下文（最多3个）
            if len(context_paragraphs) < 3:
                context_paragraphs.append(content_item)
    
    if not meaningful_content:
        return None, None, []
    
    # 按优先级排序，选择最佳内容
    meaningful_content.sort(key=lambda x: x['priority'], reverse=True)
    best_content = meaningful_content[0]
    
    # 按距离排序选择最近的3个段落作为上下文
    context_paragraphs = sorted(meaningful_content, key=lambda x: x['distance'])[:3]
    
    # 生成简洁的描述名称
    text = best_content['text']
    keywords = best_content['keywords']
    
    if keywords:
        # 如果有关键词，使用关键词
        name = keywords[0]
    else:
        # 否则从文本中提取关键部分
        # 清理文本，去除特殊字符
        clean_text = re.sub(r'[^\w\u4e00-\u9fff\s]', '', text)
        words = clean_text.split()
        
        # 选择最有意义的词汇
        if len(words) > 0:
            # 取前几个词或者限制长度
            name = ''.join(words[:2]) if len(words) > 1 else words[0]
            if len(name) > 8:
                name = name[:8]
        else:
            name = '文档内容'
    
    content_type = '标题' if best_content['is_heading'] else ('描述' if best_content['is_descriptive'] else '段落')
    return name, content_type, context_paragraphs


def check_section_content(doc, section_start, section_end, image_para_idx):
    """检查章节内容是否只有图片没有有意义的文字"""
    meaningful_text = []
    image_count = 0
    
    for para_idx in range(section_start + 1, section_end):  # 跳过标题本身
        if para_idx >= len(doc.paragraphs):
            break
            
        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        
        # 检查是否包含图片
        has_image = False
        for run in para.runs:
            if run._element.xpath('.//a:blip'):  # 检查是否包含图片
                has_image = True
                image_count += 1
                break
        
        # 如果不是图片段落且有有意义的文字（排除常见的无意义文字）
        if not has_image and text:
            # 过滤掉一些无意义的文字
            ignore_patterns = [
                '[图片:',
                '扫描件',
                '复印件', 
                '正本',
                '副本',
                '原件',
                '附件'
            ]
            
            is_meaningful = True
            for pattern in ignore_patterns:
                if pattern in text:
                    is_meaningful = False
                    break
            
            # 如果文字长度太短，也认为不是有意义的内容
            if is_meaningful and len(text) > 5:
                meaningful_text.append(text)
    
    # 如果章节内只有图片，没有有意义的文字
    return len(meaningful_text) == 0 and image_count > 0


def find_continuous_image_groups(doc, image_parts):
    """
    识别文档中的连续图片组（包括段落和表格中的图片）
    
    Args:
        doc: Word文档对象
        image_parts: 图片部件字典
        
    Returns:
        list: 连续图片组列表，每组包含图片信息和组ID
    """
    all_images = []  # 存储所有图片的位置信息
    
    # 扫描所有段落，记录图片位置
    for para_idx, paragraph in enumerate(doc.paragraphs):
        paragraph_images = []
        
        for run_idx, run in enumerate(paragraph.runs):
            if run._element.xml:
                for drawing in run._element.findall('.//w:drawing', 
                                                   {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    for blip in drawing.findall('.//a:blip', 
                                               {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        
                        if embed_id and embed_id in image_parts:
                            paragraph_images.append({
                                'para_idx': para_idx,
                                'run_idx': run_idx,
                                'embed_id': embed_id,
                                'paragraph': paragraph,
                                'location_type': 'paragraph',
                                'table_info': None
                            })
        
        all_images.extend(paragraph_images)
    
    # 扫描所有表格，记录图片位置
    table_para_offset = len(doc.paragraphs)  # 表格图片的虚拟段落索引
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for cell_para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        if run._element.xml:
                            for drawing in run._element.findall('.//w:drawing', 
                                                               {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                for blip in drawing.findall('.//a:blip', 
                                                           {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    
                                    if embed_id and embed_id in image_parts:
                                        # 为表格图片分配虚拟段落索引
                                        virtual_para_idx = table_para_offset + table_idx * 100 + row_idx * 10 + cell_idx
                                        all_images.append({
                                            'para_idx': virtual_para_idx,
                                            'run_idx': run_idx,
                                            'embed_id': embed_id,
                                            'paragraph': paragraph,
                                            'location_type': 'table',
                                            'table_info': {
                                                'table_idx': table_idx,
                                                'row_idx': row_idx,
                                                'cell_idx': cell_idx,
                                                'cell_para_idx': cell_para_idx
                                            }
                                        })
    
    # 按位置索引排序
    all_images.sort(key=lambda x: x['para_idx'])
    
    # 识别连续图片组
    image_groups = []
    current_group = []
    group_id = 0
    
    for i, img_info in enumerate(all_images):
        if not current_group:
            # 开始新组
            current_group = [img_info]
        else:
            # 判断是否属于同一连续组
            prev_img = current_group[-1]
            
            # 对于段落图片，使用原有逻辑
            if (img_info['location_type'] == 'paragraph' and 
                prev_img['location_type'] == 'paragraph'):
                
                para_gap = img_info['para_idx'] - prev_img['para_idx']
                is_continuous = True
                
                if para_gap > 0:
                    # 检查中间段落是否都是空白或只有很少文字
                    for check_para_idx in range(prev_img['para_idx'] + 1, img_info['para_idx']):
                        check_para = doc.paragraphs[check_para_idx]
                        if not is_empty_or_whitespace_paragraph(check_para):
                            if len(check_para.text.strip()) > 20:
                                is_continuous = False
                                break
                
                # 如果间隔超过3个段落，也认为不连续
                if para_gap > 3:
                    is_continuous = False
                    
            # 对于表格图片，同一表格同一行的连续单元格认为是连续的
            elif (img_info['location_type'] == 'table' and 
                  prev_img['location_type'] == 'table'):
                
                curr_table = img_info['table_info']
                prev_table = prev_img['table_info']
                
                # 同一表格同一行的相邻单元格
                is_continuous = (curr_table['table_idx'] == prev_table['table_idx'] and
                               curr_table['row_idx'] == prev_table['row_idx'] and
                               abs(curr_table['cell_idx'] - prev_table['cell_idx']) <= 2)
                
            # 表格图片和段落图片之间一般不认为连续
            else:
                is_continuous = False
            
            if is_continuous:
                # 属于当前组
                current_group.append(img_info)
            else:
                # 结束当前组，开始新组
                if len(current_group) > 1:
                    # 只有多于1张图片才算连续组
                    for img in current_group:
                        img['group_id'] = group_id
                        img['group_size'] = len(current_group)
                    group_id += 1
                else:
                    # 单张图片不算组
                    current_group[0]['group_id'] = None
                    current_group[0]['group_size'] = 1
                
                image_groups.extend(current_group)
                current_group = [img_info]
    
    # 处理最后一组
    if current_group:
        if len(current_group) > 1:
            for img in current_group:
                img['group_id'] = group_id
                img['group_size'] = len(current_group)
        else:
            current_group[0]['group_id'] = None
            current_group[0]['group_size'] = 1
        image_groups.extend(current_group)
    
    return image_groups


def convert_doc_to_docx(doc_path):
    """将.doc转换为.docx"""
    if not DOC_SUPPORT:
        raise Exception("不支持.doc格式，请先转换为.docx或安装pywin32")
    
    print(f"正在转换 {doc_path} ...")
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False
    
    try:
        doc = word_app.Documents.Open(str(Path(doc_path).absolute()))
        docx_path = str(Path(doc_path).with_suffix('.docx'))
        doc.SaveAs2(docx_path, FileFormat=16)
        doc.Close()
        word_app.Quit()
        return docx_path
    except Exception as e:
        word_app.Quit()
        raise e


def extract_and_separate(input_file):
    """
    主要功能：提取图片并分离
    
    Args:
        input_file: 输入的Word文档路径
        
    Returns:
        tuple: (分离后的文档路径, 图片文件夹路径, 图片数量)
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")
    
    # 处理.doc文件
    if input_path.suffix.lower() == '.doc':
        input_file = convert_doc_to_docx(input_file)
        input_path = Path(input_file)
    elif input_path.suffix.lower() != '.docx':
        raise ValueError(f"不支持的文件格式: {input_path.suffix}")
    
    # 创建输出路径
    base_name = input_path.stem
    output_doc = input_path.parent / f"{base_name}_分离后.docx"
    output_images = input_path.parent / f"{base_name}_图片"
    output_images.mkdir(exist_ok=True)
    
    # 加载文档
    doc = Document(input_file)
    
    # 初始化AI客户端
    ai_client = init_ai_client()
    
    # 获取所有图片
    image_parts = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_parts[rel.rId] = rel.target_part

    # 初始化AI客户端
    ai_client = init_ai_client()
    
    # 识别连续图片组
    image_groups = find_continuous_image_groups(doc, image_parts)
    
    image_count = 0
    processed_images = set()  # 记录已处理的图片，避免重复处理
    group_base_names = {}  # 缓存每个组的基础名称

    # 按连续组处理图片
    for img_info in image_groups:
        if img_info['embed_id'] in processed_images:
            continue
            
        image_count += 1
        processed_images.add(img_info['embed_id'])
        
        # 收集上下文 - 改进版本
        context_text = ""
        
        # 首先确定图片在组中的位置
        if img_info['group_id'] is not None:
            # 是连续组的一部分，计算在组中的位置
            group_images = [img for img in image_groups if img['group_id'] == img_info['group_id']]
            group_images.sort(key=lambda x: x['para_idx'])
            img_index_in_group = group_images.index(img_info) + 1
            total_in_group = len(group_images)
            
            # 检查是否已经为这个组生成了基础名称
            if img_info['group_id'] in group_base_names:
                # 使用已生成的基础名称
                base_name = group_base_names[img_info['group_id']]
                naming_method = "连续组共享命名"
            else:
                # 这是组中的第一张图片，需要生成基础名称
                base_name = None
                naming_method = ""
        else:
            # 单独图片
            img_index_in_group = 1
            total_in_group = 1
            base_name = None
            naming_method = ""
        
        # 首先收集上下文信息用于AI命名
        # 1. 获取文档标题
        doc_title = ""
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if first_para:
                doc_title = first_para[:50]  # 取前50个字符作为文档标题
        
        # 2. 为每张图片寻找最近的5段有效上下文文字
        nearest_contexts = find_nearest_context_paragraphs(doc, img_info['para_idx'])
        
        # 3. 组合上下文信息（优先使用最近的5段上下文）
        context_parts = []
        
        # 添加文档标题
        if doc_title:
            context_parts.append(f"文档标题: {doc_title}")
        
        # 添加最近的5段上下文（这是最重要的信息）
        if nearest_contexts:
            context_parts.append("最有效上下文:")
            for i, ctx in enumerate(nearest_contexts, 1):
                # 限制每段上下文的长度以避免提示词过长
                ctx_text = ctx['text'][:100] if len(ctx['text']) > 100 else ctx['text']
                value_info = f"评分{ctx['value_score']}"
                if ctx['has_key_info']:
                    value_info += ",关键信息"
                if ctx['is_heading']:
                    value_info += ",标题"
                context_parts.append(f"  {i}. 段落{ctx['para_idx']}(距离{ctx['distance']},{value_info}): {ctx_text}")
        
        # 添加额外的角色关键词（如果空间允许）
        role_keywords = ["授权代理人", "授权委托人", "法定代表人", "单位负责人", "自然人", "身份证", "营业执照", "执照"]
        role_context = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if para_text:
                for keyword in role_keywords:
                    if keyword in para_text:
                        role_context.append(f"{keyword}: {para_text[:60]}")  # 缩短长度
                        break
                        
        # 只添加最相关的角色上下文（最多2个）
        if role_context:
            context_parts.append("角色信息:")
            for role_ctx in role_context[:2]:
                context_parts.append(f"  - {role_ctx}")
            
        context_text = "\n".join(context_parts)
        
        # 优先AI命名，AI不可用或AI命名不合理时fallback为标题名
        ai_available = ai_client is not None
        if not base_name and ai_available and context_text.strip():
            try:
                if os.getenv('AI_NAME_LOG'):
                    print(f"🤖 准备调用AI命名...")
                    print(f"   图片段落: {img_info['para_idx']}")
                    print(f"   组ID: {img_info['group_id']}")
                    print(f"   上下文长度: {len(context_text)} 字符")
                    print(f"   上下文预览: {context_text[:200]}...")
                ai_name = generate_image_name(
                    ai_client, 
                    context_text.strip(),
                    img_index_in_group,
                    total_in_group
                )
                if os.getenv('AI_NAME_LOG'):
                    print(f"🎯 AI返回结果: '{ai_name}'")
                # 从AI命名中提取基础名称
                base_name = re.sub(r'\d+$', '', ai_name)  # 移除末尾数字
                if not base_name:
                    base_name = ai_name
                # 校验AI命名是否合理
                if _is_ai_result_valid(base_name, context_text):
                    naming_method = "AI智能命名"
                    if img_info['group_id'] is not None:
                        group_base_names[img_info['group_id']] = base_name
                        if os.getenv('AI_NAME_LOG'):
                            print(f"💾 为组 {img_info['group_id']} 缓存基础名称: '{base_name}'")
                else:
                    if os.getenv('AI_NAME_LOG'):
                        print(f"⚠️ AI命名不合理: '{base_name}'，使用fallback")
                    base_name = None
            except Exception as e:
                if os.getenv('AI_NAME_LOG'):
                    print(f"⚠️ AI命名失败: {e}")
                ai_available = False
                base_name = None
        
        # 如果AI命名失败或不可用，直接选取图片所属的标题名
        if not base_name and nearest_contexts:
            # 优先使用最近的标题
            for context in nearest_contexts:
                if context.get('is_heading') and context['text']:
                    title = context['text'].strip()
                    clean_title = re.sub(r'[^\w\u4e00-\u9fff]', '', title)
                    if 2 <= len(clean_title) <= 15:
                        base_name = clean_title
                        naming_method = "标题命名"
                        break
            # 如果没有合适的标题，使用第一个高价值上下文

            if not base_name:
                first_context = nearest_contexts[0]
                if first_context.get('has_key_info') and first_context['text']:
                    text = first_context['text'].strip()
                    clean_text = re.sub(r'[^\w\u4e00-\u9fff]', '', text)
                    if 2 <= len(clean_text) <= 15:
                        base_name = clean_text
                        naming_method = "上下文命名"


# --- 保证顶格定义 ---
def _is_ai_result_valid(ai_name, context_text):
    """
    校验AI返回的命名是否合理
    """
    if not ai_name or len(ai_name.strip()) == 0:
        return False
    ai_name = ai_name.strip()
    # 如果AI命名为营业执照，但上下文没有营业执照，则判为不合理
    if '营业执照' in ai_name and '营业执照' not in context_text:
        return False
    # 其它简单规则：命名长度合理且不是纯数字
    if 2 <= len(ai_name) <= 12 and not ai_name.isdigit():
        return True
    return False

    # 如果最近上下文也没有，检查章节标题命名
if not base_name:
    headings = collect_headings(doc)
    section_info, section_start, section_end = find_section_for_image(headings, img_info['para_idx'], len(doc.paragraphs))
    if section_info and check_section_content(doc, section_start, section_end, img_info['para_idx']):
        section_title = section_info['text']
        base_name = re.sub(r'[^\w\u4e00-\u9fff]', '', section_title)
        if len(base_name) > 12:
            base_name = base_name[:12]
        naming_method = "章节标题命名"

# 最后的兜底方案
if not base_name:
    base_name = "图片"
    naming_method = "默认命名"
        
        # 调试日志
        if os.getenv('AI_NAME_LOG'):
            print(f"🤖 使用{naming_method}: '{base_name}'")
            if nearest_contexts:
                print(f"🔍 找到的5段最有效上下文:")
                for i, ctx in enumerate(nearest_contexts, 1):
                    ctx_type_info = []
                    if ctx['has_key_info']:
                        ctx_type_info.append("关键信息")
                    if ctx['is_heading']:
                        ctx_type_info.append("标题")
                    ctx_type = ",".join(ctx_type_info) if ctx_type_info else "普通"
                    print(f"   {i}. 段落{ctx['para_idx']}(距离{ctx['distance']},评分{ctx['value_score']},{ctx_type}): {ctx['text'][:50]}...")
        
        # 使用新的命名格式：图片顺序+图片名称+（连续图片序号）
        if total_in_group > 1:
            final_name = f"{image_count}{base_name}（{img_index_in_group}）"
        else:
            final_name = f"{image_count}{base_name}"
        
        # 调试日志
        if os.getenv('AI_NAME_LOG'):
            print(f"🏷️  最终命名: '{final_name}' (基础名: {base_name}, 图片序号: {image_count}, 组内序号: {img_index_in_group}/{total_in_group})")
        
        # 保存图片
        image_filename = save_image(
            image_parts[img_info['embed_id']], 
            output_images, 
            image_count,
            custom_name=final_name
        )
        
        # 在文档中替换图片为文本标记
        paragraph = img_info['paragraph']
        run_idx = img_info['run_idx']
        
        # 移除包含图片的run
        if run_idx < len(paragraph.runs):
            paragraph._element.remove(paragraph.runs[run_idx]._element)
        
        # 添加替换文本
        replacement_text = f"[图片: {image_filename}]"
        if paragraph.runs:
            # 如果段落还有其他runs，添加到第一个run
            if paragraph.runs[0].text:
                paragraph.runs[0].text = replacement_text + " " + paragraph.runs[0].text
            else:
                paragraph.runs[0].text = replacement_text
        else:
            # 如果段落没有runs，创建新的run
            paragraph.add_run(replacement_text)

    # 保存分离后的文档
    doc.save(str(output_doc))

    return str(output_doc), str(output_images), image_count


def save_image(image_part, output_dir, counter, custom_name=None):
    """保存图片文件"""
    image_data = image_part.blob
    
    # 检测图片格式
    try:
        img = Image.open(BytesIO(image_data))
        format_name = img.format.lower() if img.format else 'png'
        if format_name == 'jpeg':
            format_name = 'jpg'
    except:
        format_name = 'png'
    
    # 确定文件名
    if custom_name:
        # 使用自定义名称
        filename = f"{custom_name}.{format_name}"
    else:
        # 使用默认命名
        filename = f"图片_{counter:03d}.{format_name}"
    
    # 处理文件名冲突
    filepath = output_dir / filename
    duplicate_counter = 1
    while filepath.exists():
        if custom_name:
            filename = f"{custom_name}_{duplicate_counter}.{format_name}"
        else:
            filename = f"图片_{counter:03d}_{duplicate_counter}.{format_name}"
        filepath = output_dir / filename
        duplicate_counter += 1
    
    # 保存图片
    with open(filepath, 'wb') as f:
        f.write(image_data)
    
    return filename


def main():
    """主函数"""
    print("=" * 50)
    print("Word文档图片分离工具")
    print("=" * 50)
    
    # 检查帮助
    if len(sys.argv) > 1 and sys.argv[1] in ['-h', '--help', 'help']:
        print("使用方法:")
        print(f"  python {Path(__file__).name} <Word文档路径>")
        print("\n示例:")
        print(f"  python {Path(__file__).name} document.docx")
        print(f"  python {Path(__file__).name} '包含空格的文档.docx'")
        print("\n输出:")
        print("  - 文档名_分离后.docx (图文分离后的文档)")
        print("  - 文档名_图片/ (包含所有图片的文件夹)")
        return
    
    # 获取输入文件
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        input_file = input("请输入Word文档路径: ").strip('"')
    
    if not input_file:
        print("错误: 未指定输入文件")
        print("使用 --help 查看帮助")
        return
    
    try:
        # 执行分离
        print(f"正在处理: {Path(input_file).name}")
        output_doc, output_images, image_count = extract_and_separate(input_file)
        # 显示结果
        print("\n处理完成!")
        print(f"分离后文档: {Path(output_doc).name}")
        print(f"图片文件夹: {Path(output_images).name}")
        print(f"提取图片数量: {image_count}")
        if image_count == 0:
            print("文档中没有找到图片")
    except Exception as e:
        print(f"处理失败: {e}")


if __name__ == "__main__":
    main()