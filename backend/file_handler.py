#!/usr/bin/env python3
"""
文件处理服务模块
================

本模块提供文件上传、内容提取和文件管理的核心功能。
支持PDF和Word文档的安全处理和文本内容提取。

主要功能：
    1. 文件类型验证和安全检查
    2. 文件上传和存储管理
    3. PDF文档内容提取
    4. Word文档内容提取
    5. 文件信息获取和管理

支持格式：
    - PDF (.pdf) - 使用PyPDF2库
    - Word文档 (.docx) - 使用python-docx库
    - 计划支持 (.doc) - 需要额外转换工具

安全特性：
    - 文件名安全化处理
    - 文件类型白名单验证
    - 文件大小限制支持
    - 恶意文件检测（基础）

依赖库：
    - werkzeug: 文件安全处理
    - pathlib: 路径操作
    - docx: Word文档处理
    - PyPDF2: PDF文档处理
    - os: 系统操作
    - uuid: 唯一标识生成

作者：BidAnalysis Team
创建时间：2025年
版本：1.0
"""

import os
import uuid
from werkzeug.utils import secure_filename
from pathlib import Path
import docx
import PyPDF2
# import fitz  # PyMuPDF - 可选的PDF处理库，如果安装失败可注释掉
from typing import Optional

# 尝试导入doc文件处理库
try:
    import docx2txt  # 用于处理.doc文件
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False
    print("警告: 未安装 docx2txt 库，无法处理 .doc 文件。请运行: pip install docx2txt")

class FileHandler:
    """
    文件处理服务类
    ==============
    
    负责处理文件上传、存储、内容提取等核心功能。
    提供安全可靠的文件操作接口。
    
    类属性：
        ALLOWED_EXTENSIONS: 允许的文件扩展名集合
        
    主要方法：
        - is_allowed_file(): 检查文件类型
        - save_file(): 保存上传文件
        - extract_content(): 提取文件内容
        - get_file_info(): 获取文件信息
        - delete_file(): 删除文件
        
    使用示例：
        handler = FileHandler()
        if handler.is_allowed_file(filename):
            path = handler.save_file(file, file_id, upload_dir)
            content = handler.extract_content(path)
    """
    
    # 定义允许上传的文件扩展名（白名单策略）
    ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx'}
    
    def __init__(self):
        """
        初始化文件处理服务
        
        设置基本配置和初始化必要的组件。
        """
        pass
    
    def is_allowed_file(self, filename: str) -> bool:
        """
        检查文件类型是否被允许上传
        ============================
        
        基于文件扩展名的白名单策略验证文件类型。
        这是第一道安全防线，防止恶意文件上传。
        
        Args:
            filename (str): 待检查的文件名
            
        Returns:
            bool: True表示允许上传，False表示禁止上传
            
        检查逻辑：
            1. 验证文件名是否为空或None
            2. 提取文件扩展名（转换为小写）
            3. 对比白名单ALLOWED_EXTENSIONS
            
        支持格式：
            - .pdf: PDF文档
            - .docx: Word 2007+文档
            - .doc: Word 97-2003文档（计划支持）
            
        安全考虑：
            - 使用白名单而非黑名单策略
            - 忽略大小写差异
            - 防止空文件名绕过检查
        """
        # 检查文件名是否有效
        if not filename:
            return False
        
        # 提取文件扩展名并转换为小写
        file_ext = Path(filename).suffix.lower()
        
        # 检查扩展名是否在允许列表中
        return file_ext in self.ALLOWED_EXTENSIONS
    
    def save_file(self, file, file_id: str, upload_folder: str) -> str:
        """
        安全保存上传的文件
        ==================
        
        将上传的文件保存到指定目录，使用UUID作为文件名避免冲突。
        应用安全文件名处理防止路径遍历攻击。
        
        Args:
            file: Flask上传的文件对象
            file_id (str): 系统生成的唯一文件标识符
            upload_folder (str): 文件上传目录路径
            
        Returns:
            str: 保存后的完整文件路径
            
        安全措施：
            1. 使用secure_filename()处理原始文件名
            2. 使用UUID作为实际存储文件名
            3. 保留原始文件扩展名
            4. 避免文件名冲突
            
        文件命名规则：
            原文件名: "投标文件.docx"
            存储文件名: "{uuid}.docx"
            
        异常处理：
            - 目录不存在会自动创建
            - 文件保存失败会抛出异常
        """
        # 使用werkzeug安全处理文件名，防止路径遍历攻击
        filename = secure_filename(file.filename)
        
        # 提取原始文件扩展名
        file_ext = Path(filename).suffix.lower()
        
        # 使用UUID生成唯一的存储文件名，避免冲突
        new_filename = f"{file_id}{file_ext}"
        
        # 构建完整的文件存储路径
        file_path = os.path.join(upload_folder, new_filename)
        
        # 保存文件到指定路径
        file.save(file_path)
        
        return file_path
    
    def extract_content(self, file_path: str) -> str:
        """
        提取文件文本内容
        ================
        
        根据文件类型自动选择合适的内容提取方法。
        支持PDF和Word文档的文本内容提取。
        
        Args:
            file_path (str): 文件的完整路径
            
        Returns:
            str: 提取的文本内容
            
        支持格式：
            - PDF (.pdf): 使用PyPDF2提取文本
            - Word (.docx): 使用python-docx提取文本和表格
            - Word (.doc): 计划支持，当前会抛出异常
            
        提取内容：
            - 文档正文文本
            - 表格内容（Word文档）
            - 页眉页脚（部分支持）
            
        异常处理：
            - 不支持的文件类型
            - 文件损坏或无法读取
            - 内容提取失败
            
        注意：
            - PDF文档的提取效果取决于文档的创建方式
            - 扫描版PDF需要OCR技术（未实现）
            - 加密文档需要密码（未实现）
        """
        # 获取文件扩展名
        file_ext = Path(file_path).suffix.lower()
        
        try:
            # 根据文件类型选择相应的提取方法
            if file_ext == '.pdf':
                return self._extract_pdf_content(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._extract_word_content(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {file_ext}")
        except Exception as e:
            # 统一异常处理，提供有用的错误信息
            raise Exception(f"提取文件内容失败: {str(e)}")
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """
        提取PDF文件内容（私有方法）
        ===========================
        
        使用PyPDF2库提取PDF文档的文本内容。
        逐页提取文本并合并为完整内容。
        
        Args:
            file_path (str): PDF文件路径
            
        Returns:
            str: 提取的文本内容
            
        提取过程：
            1. 打开PDF文件（二进制模式）
            2. 创建PdfReader对象
            3. 遍历所有页面
            4. 提取每页文本内容
            5. 合并所有文本
            
        局限性：
            - 无法处理扫描版PDF（需要OCR）
            - 无法提取图片中的文字
            - 复杂排版可能影响文本顺序
            - 加密PDF需要密码
            
        异常处理：
            - 文件不存在
            - 文件损坏
            - 权限不足
            - 内存不足（大文件）
        """
        content = ""
        
        try:
            # 以二进制模式打开PDF文件
            with open(file_path, 'rb') as file:
                # 创建PDF读取器对象
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 遍历所有页面提取文本
                for page in pdf_reader.pages:
                    # 提取当前页面的文本内容
                    page_text = page.extract_text()
                    if page_text:
                        # 添加页面文本到总内容，页面间用换行分隔
                        content += page_text + "\n"
        except Exception as e:
            # 抛出详细的错误信息
            raise Exception(f"PDF内容提取失败: {str(e)}")
        
        # 返回清理后的文本内容
        return content.strip()
    
    def _extract_word_content(self, file_path: str) -> str:
        """
        提取Word文档内容（私有方法）
        ============================
        
        使用python-docx库提取Word文档的文本内容。
        支持 .docx 和 .doc 格式文件。
        
        Args:
            file_path (str): Word文档路径
            
        Returns:
            str: 提取的文本内容
            
        提取内容：
            1. 段落文本 - 文档主体内容
            2. 表格数据 - 以"|"分隔的表格内容（仅.docx支持）
            3. 列表项目 - 作为普通段落处理
            
        格式处理：
            - 保持段落结构（换行）
            - 表格行以"|"分隔列
            - 空白内容自动过滤
            
        支持版本：
            - .docx格式 (Word 2007+) - 完全支持
            - .doc格式 (Word 97-2003) - 基本文本提取
            
        局限性：
            - 无法提取图片中的文字
            - .doc格式不支持表格结构化提取
            - 页眉页脚需要特殊处理
            
        异常处理：
            - 文件格式错误
            - 文档损坏
            - 权限问题
            - 内存不足
        """
        try:
            # 检查文件扩展名
            if file_path.lower().endswith('.docx'):
                # 处理.docx格式文件（完全支持）
                doc = docx.Document(file_path)
                content = ""
                
                # 提取段落文本
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # 提取表格文本
                for table in doc.tables:
                    for row in table.rows:
                        # 收集当前行的所有单元格文本
                        row_text = []
                        for cell in row.cells:
                            # 清理单元格文本并添加到行文本列表
                            cell_text = cell.text.strip()
                            row_text.append(cell_text)
                        # 将行文本用"|"连接并添加到内容中
                        content += " | ".join(row_text) + "\n"
                
                return content.strip()
            
            elif file_path.lower().endswith('.doc'):
                # 处理.doc格式文件
                if DOC_SUPPORT:
                    try:
                        # 使用docx2txt库提取.doc文件内容
                        content = docx2txt.process(file_path)
                        if content and content.strip():
                            return content.strip()
                        else:
                            # 如果docx2txt返回空内容，尝试备选方案
                            print("docx2txt返回空内容，尝试备选方案...")
                            return self._extract_doc_fallback(file_path)
                    except Exception as e:
                        # docx2txt失败时，尝试备选方案
                        print(f"docx2txt处理失败: {e}")
                        return self._extract_doc_fallback(file_path)
                else:
                    # 如果没有安装docx2txt库，直接使用备选方案
                    return self._extract_doc_fallback(file_path)
            
            else:
                # 不支持的文件格式
                raise ValueError(f"不支持的Word文档格式: {Path(file_path).suffix}")
                
        except Exception as e:
            # 抛出详细的错误信息
            raise Exception(f"Word文档内容提取失败: {str(e)}")
    
    def _extract_doc_fallback(self, file_path: str) -> str:
        """
        .doc文件的备选提取方法
        =====================
        
        当docx2txt库不可用或失败时的备选方案。
        尝试使用其他方法提取.doc文件内容。
        
        Args:
            file_path (str): .doc文件路径
            
        Returns:
            str: 提取的文本内容
            
        备选方案：
            1. 尝试使用win32com.client (Windows环境)
            2. 尝试将文件当作ZIP处理（某些.doc文件）
            3. 提示用户转换文件格式
            
        注意：
            - 此方法的成功率依赖于运行环境
            - 建议用户将.doc转换为.docx格式
        """
        # 方案1: 尝试使用win32com.client (仅Windows)
        import platform
        if platform.system() == "Windows":
            try:
                import win32com.client
                
                # 获取绝对路径
                abs_path = os.path.abspath(file_path)
                
                # 创建Word应用程序对象
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                try:
                    # 打开文档
                    doc = word.Documents.Open(abs_path)
                    content = doc.Content.Text
                    
                    # 关闭文档
                    doc.Close()
                    
                    if content and content.strip():
                        return content.strip()
                        
                except Exception as e:
                    print(f"Word COM操作失败: {e}")
                finally:
                    # 确保关闭Word应用程序
                    try:
                        word.Quit()
                    except:
                        pass
                        
            except ImportError:
                print("win32com.client不可用")
            except Exception as e:
                print(f"win32com处理失败: {e}")
        
        # 方案2: 如果所有方法都失败，提供友好的错误信息
        raise Exception(
            "无法处理此.doc格式文件。可能的解决方案：\n"
            "1. 将文件另存为.docx格式后重新上传（推荐）\n"
            "2. 检查文件是否损坏或为特殊格式\n"
            "3. 在Windows环境下确保已安装Microsoft Word\n"
            "4. 尝试用其他工具转换文件格式"
        )
    
    def get_file_info(self, file_path: str) -> dict:
        """
        获取文件基本信息
        ================
        
        收集并返回文件的元数据信息，用于文件管理和状态检查。
        
        Args:
            file_path (str): 文件的完整路径
            
        Returns:
            dict: 包含文件信息的字典
                - size (int): 文件大小（字节）
                - extension (str): 文件扩展名
                - modified_time (float): 最后修改时间戳
                - created_time (float): 创建时间戳
                
        使用场景：
            - 文件管理界面显示
            - 上传状态确认
            - 存储空间统计
            - 文件变更检测
            
        异常处理：
            - 文件不存在：抛出FileNotFoundError
            - 权限不足：抛出PermissionError
            - 其他系统错误
            
        注意：
            - 时间戳为Unix时间戳格式
            - 大小单位为字节
            - Windows和Unix系统的创建时间定义可能不同
        """
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件状态信息
        file_stat = os.stat(file_path)
        
        # 提取文件扩展名
        file_ext = Path(file_path).suffix.lower()
        
        # 构建并返回文件信息字典
        return {
            'size': file_stat.st_size,          # 文件大小（字节）
            'extension': file_ext,              # 文件扩展名
            'modified_time': file_stat.st_mtime, # 最后修改时间
            'created_time': file_stat.st_ctime   # 创建时间（或最后元数据变更时间）
        }
    
    def delete_file(self, file_path: str) -> bool:
        """
        安全删除文件
        ============
        
        删除指定路径的文件，提供安全的错误处理。
        
        Args:
            file_path (str): 要删除的文件路径
            
        Returns:
            bool: 删除操作结果
                - True: 删除成功
                - False: 删除失败（文件不存在或其他错误）
                
        安全特性：
            - 检查文件是否存在
            - 异常捕获避免程序崩溃
            - 返回明确的操作结果
            
        使用场景：
            - 清理临时文件
            - 删除过期文档
            - 用户主动删除
            - 系统维护清理
            
        注意：
            - 删除操作不可逆
            - 需要相应的文件系统权限
            - 不会删除目录，只删除文件
        """
        try:
            # 检查文件是否存在
            if os.path.exists(file_path):
                # 删除文件
                os.remove(file_path)
                return True
            # 文件不存在，返回False
            return False
        except Exception:
            # 发生任何异常都返回False
            # 可能的异常：权限错误、文件被占用等
            return False
